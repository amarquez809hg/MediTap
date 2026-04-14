#!/usr/bin/env python3
"""Convert docs/AGENT_SESSION_CHANGELOG.md to a Word .docx (requires python-docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
except ImportError as e:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e

ROOT = Path(__file__).resolve().parents[2]
MD_PATH = ROOT / "docs" / "AGENT_SESSION_CHANGELOG.md"
OUT_PATH = ROOT / "docs" / "AGENT_SESSION_CHANGELOG.docx"


def is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|").replace(" ", "")
    return "---" in inner and set(inner) <= {"|", "-", ":", " "}


def split_table_row(line: str) -> list[str] | None:
    line = line.rstrip("\n")
    if not line.strip().startswith("|"):
        return None
    parts = [p.strip() for p in line.strip().split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts if parts else None


def add_runs(p, text: str) -> None:
    """Inline **bold** and `code` (non-nested)."""
    if not text:
        return
    pat = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        chunk = m.group(1)
        if chunk.startswith("**"):
            p.add_run(chunk[2:-2]).bold = True
        else:
            r = p.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
    doc.add_paragraph()


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    cp = doc.core_properties
    cp.title = "MediTap — Agent session change register"
    cp.subject = "Jira-style changelog (Sets 1–4)"
    cp.keywords = "MediTap;changelog;agent"

    table_buf: list[list[str]] = []

    def end_table() -> None:
        nonlocal table_buf
        if table_buf:
            flush_table(doc, table_buf)
            table_buf = []

    for line in lines:
        raw = line.rstrip()

        if raw.strip().startswith("|"):
            if is_table_sep(raw):
                continue
            row = split_table_row(raw)
            if row is not None:
                table_buf.append(row)
            continue
        else:
            end_table()

        if raw.strip() == "---":
            doc.add_paragraph("—" * 20).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue

        if raw.startswith("# "):
            doc.add_heading(raw[2:].strip(), level=1)
            continue
        if raw.startswith("## "):
            doc.add_heading(raw[3:].strip(), level=2)
            continue
        if raw.startswith("### "):
            doc.add_heading(raw[4:].strip(), level=3)
            continue
        if raw.startswith("#### "):
            doc.add_heading(raw[5:].strip(), level=4)
            continue

        if raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, raw[2:].strip())
            continue

        if re.match(r"^\d+\.\s+", raw):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", raw).strip())
            continue

        if not raw.strip():
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        add_runs(p, raw.strip())

    end_table()

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
